import argparse, json, os, sys
from pathlib import Path
from datetime import datetime

SCHEMA_VERSION = 1

def find_project(start):
    """只读查找项目根目录，找不到返回 None，不创建任何文件"""
    c = Path(start).resolve()
    state_file = c / ".novel" / "state.json"
    if state_file.exists():
        return c
    for child in sorted(c.iterdir()):
        if child.is_dir():
            child_state = child / ".novel" / "state.json"
            if child_state.exists():
                return child
    return None

def migrate_state(state):
    """迁移旧版 state.json 到当前 schema。

    历史版本：
      - 无 schema_version 字段：早期版本，可能用 'project_info' 而非 'project'
      - schema_version=1：当前版本，统一用 'project'

    返回 (migrated_state, changed) 元组。changed=True 表示发生了迁移需要持久化。
    幂等，可重复调用。
    """
    if not isinstance(state, dict):
        return state, False
    # 已是当前版本
    if state.get("schema_version") == SCHEMA_VERSION:
        return state, False
    # 旧版迁移：project_info -> project
    if "project_info" in state and "project" not in state:
        state["project"] = state.pop("project_info")
    state["schema_version"] = SCHEMA_VERSION
    return state, True

def init_project(start, title="", author="", genre=""):
    """显式创建新项目，返回项目根"""
    c = Path(start).resolve()
    novel_dir = c / ".novel"
    novel_dir.mkdir(parents=True, exist_ok=True)
    state_file = novel_dir / "state.json"
    init_state = {
        "schema_version": SCHEMA_VERSION,
        "project": {"book_name": title, "author": author, "genre": genre},
        "progress": {},
        "versions": {"baseline_version": 1}
    }
    state_file.write_text(json.dumps(init_state, ensure_ascii=False, indent=2), "utf-8")
    return c

def load_state(r):
    f = r / ".novel" / "state.json"
    if f.exists():
        try:
            state = json.loads(f.read_text("utf-8"))
            # 自动迁移旧版 schema（含 project_info 兼容），持久化迁移结果
            migrated, changed = migrate_state(state)
            if changed:
                save_state(r, migrated)
            return migrated
        except Exception:
            return {}
    return {}

def save_state(r, s):
    f = r / ".novel" / "state.json"
    f.parent.mkdir(parents=True, exist_ok=True)
    f.write_text(json.dumps(s, ensure_ascii=False, indent=2), "utf-8")

def cmd_init(a):
    r = init_project(Path(a.project_root), a.title or "", a.author or "", a.genre or "")
    print(str(r))

def cmd_where(a):
    r = find_project(Path(a.project_root))
    if r:
        print(str(r))
    else:
        print("no project found", file=sys.stderr)
        sys.exit(1)

def cmd_status(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    s = load_state(r)
    if not s:
        print("no state")
        return
    # load_state 已自动迁移 'project_info' -> 'project'，这里直接用 'project'
    pi = s.get("project", {})
    pr = s.get("progress", {})
    print("book:", pi.get("book_name", pi.get("title", "?")))
    print("genre:", pi.get("genre", "?"))
    print("writing_started:", pr.get("writing_started", False))

def cmd_doctor(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    issues = []
    try:
        checks = [
            (r / ".novel" / "state.json", "state.json"),
            (r / "设定集" / "主角卡.md", "设定集/主角卡.md"),
            (r / "设定集" / "世界观.md", "设定集/世界观.md"),
            (r / "大纲" / "总纲.md", "大纲/总纲.md"),
        ]
        for path_obj, name in checks:
            if not path_obj.exists():
                issues.append("missing: " + name)
        td = r / "正文"
        chs = []
        if td.exists():
            chs = sorted(td.glob("第*.md"))
            if chs:
                print("chapters:", len(chs))
        tracks = ["追踪/上下文.md", "追踪/伏笔.md", "追踪/时间线.md", "追踪/角色状态.md"]
        for t in tracks:
            if not (r / t).exists():
                issues.append("missing: " + t)

        chapter_arg = getattr(a, "chapter", None)
        if chapter_arg:
            try:
                ch_num = int(chapter_arg)
                ch_file = None
                for ch in chs:
                    if ("第" + str(ch_num).zfill(2)) in ch.name or ("第" + str(ch_num)) in ch.name:
                        ch_file = ch
                        break
                if ch_file:
                    print("chapter", ch_num, ":", ch_file.name)
                    if getattr(a, "deep", False):
                        body = ch_file.read_text("utf-8")
                        char_count = len(body)
                        print("  chars:", char_count)
                        if char_count < 1500:
                            issues.append("chapter " + str(ch_num) + " too short: " + str(char_count) + " chars (< 1500)")
                        review_file = r / "审查报告" / ("第" + str(ch_num).zfill(2) + "章审查报告.md")
                        if not review_file.exists():
                            issues.append("missing review for chapter " + str(ch_num))
                        commit_file = r / ".story-system" / "commits" / ("chapter_" + str(ch_num).zfill(4) + ".commit.json")
                        if not commit_file.exists():
                            issues.append("missing commit for chapter " + str(ch_num))
                else:
                    issues.append("chapter " + str(ch_num) + " not found in 正文/")
            except ValueError:
                issues.append("invalid chapter number: " + str(chapter_arg))

        if getattr(a, "deep", False) and not chapter_arg:
            ss = r / ".story-system"
            if not ss.exists():
                issues.append("missing: .story-system/ (底本目录)")
            else:
                master = ss / "MASTER_SETTING.json"
                if not master.exists():
                    issues.append("missing: .story-system/MASTER_SETTING.json")

        if a.format == "json":
            print(json.dumps({"issues": issues, "ok": len(issues) == 0}, ensure_ascii=False))
        else:
            for i in issues:
                print(i)
            if not issues:
                print("project structure OK")
    except Exception as e:
        print("diagnose error: " + str(e))
        print("tip: check if project path contains special characters")

def cmd_contract(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    s = load_state(r)
    # load_state 已自动迁移，直接用 'project'
    pi = s.get("project", {})
    d = r / ".story-system"
    d.mkdir(parents=True, exist_ok=True)
    m = {"route": {"primary_genre": pi.get("genre", ""), "target_platform": "fanqie"}, "versions": {"baseline_version": 1}}
    (d / "MASTER_SETTING.json").write_text(json.dumps(m, ensure_ascii=False, indent=2), "utf-8")
    print("baseline done")

def cmd_export(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    td = r / "正文"
    chs = sorted(td.glob("第*.md")) if td.exists() else []
    if not chs:
        print("no chapters")
        return

    fmt = getattr(a, "format", "txt") or "txt"
    output_arg = getattr(a, "output", None)

    if output_arg:
        out_path = Path(output_arg)
        if out_path.is_dir() or output_arg.endswith(("/", "\\")):
            suffix = "txt" if fmt in ("txt", "fanqie") else "docx"
            out_path = out_path / (r.name + "_完本." + suffix)
    else:
        out_dir = r / "导出"
        out_dir.mkdir(exist_ok=True)
        suffix = "txt" if fmt in ("txt", "fanqie") else "docx"
        out_path = out_dir / (r.name + "_完本." + suffix)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    parts = [ch.read_text("utf-8") for ch in chs]

    if fmt == "fanqie":
        merged = []
        for ch in chs:
            title = ch.stem
            body = ch.read_text("utf-8").strip()
            if not title.startswith("第"):
                title = "第" + title + "章"
            merged.append(title + "\n\n" + body + "\n\n\n")
        out_path.write_text("".join(merged), encoding="utf-8")
    elif fmt == "docx":
        try:
            from docx import Document
            doc = Document()
            for ch in chs:
                title = ch.stem
                body = ch.read_text("utf-8").strip()
                doc.add_heading(title, level=1)
                for para in body.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
                doc.add_page_break()
            doc.save(str(out_path))
        except ImportError:
            fallback = out_path.with_suffix(".txt")
            fallback.write_text("\n\n".join(parts), encoding="utf-8")
            print("warning: python-docx not installed, fallback to txt:", fallback)
            out_path = fallback
    else:
        out_path.write_text("\n\n".join(parts), encoding="utf-8")

    print("exported:", out_path)
    print("  chapters:", len(chs))
    print("  total chars:", sum(len(p) for p in parts))

def cmd_commit(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    ch = int(a.chapter)
    d = r / ".story-system" / "commits"
    d.mkdir(parents=True, exist_ok=True)
    c = {"chapter": ch, "timestamp": datetime.now().isoformat(), "status": "accepted"}
    (d / f"chapter_{ch:04d}.commit.json").write_text(json.dumps(c, ensure_ascii=False, indent=2), "utf-8")
    st = load_state(r)
    st.setdefault("progress", {})["current_chapter"] = ch
    save_state(r, st)

def cmd_review(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    rf = Path(a.report_file)
    rf.parent.mkdir(parents=True, exist_ok=True)
    results = None
    if a.review_results:
        try:
            results = json.loads(a.review_results)
            if isinstance(results, str):
                results = json.loads(results)
        except Exception:
            results = a.review_results
    report_lines = [
        "# 第 " + str(a.chapter) + " 章审查报告",
        "",
        "生成: " + datetime.now().strftime("%Y-%m-%d %H:%M"),
        ""
    ]
    if results:
        if isinstance(results, dict):
            for k, v in results.items():
                report_lines.append("## " + str(k))
                report_lines.append("")
                if isinstance(v, str):
                    report_lines.append(v)
                elif isinstance(v, list):
                    for item in v:
                        report_lines.append("- " + str(item))
                else:
                    report_lines.append(str(v))
                report_lines.append("")
        elif isinstance(results, list):
            report_lines.append("## 审查结果")
            report_lines.append("")
            for idx, item in enumerate(results, 1):
                if isinstance(item, dict):
                    report_lines.append("### 条目 " + str(idx))
                    for k, v in item.items():
                        if isinstance(v, list):
                            for sub in v:
                                report_lines.append("- " + str(k) + ": " + str(sub))
                        else:
                            report_lines.append("- " + str(k) + ": " + str(v))
                    report_lines.append("")
                else:
                    report_lines.append(str(idx) + ". " + str(item))
            report_lines.append("")
        else:
            report_lines.append(str(results))
            report_lines.append("")
    else:
        report_lines.append("由 reviewer agent 输出。")
    rf.write_text("\n".join(report_lines), "utf-8")
    print("审查报告:", rf)
    if a.save_metrics and a.metrics_out:
        mp = Path(a.metrics_out)
        mp.parent.mkdir(parents=True, exist_ok=True)
        metrics = {"chapter": int(a.chapter), "timestamp": datetime.now().isoformat(), "results": results if isinstance(results, dict) else {}}
        mp.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), "utf-8")
        print("metrics saved:", mp)

def cmd_memory(a):
    r = find_project(Path(a.project_root))
    if not r:
        print("no project")
        return
    mf = r / ".novel" / "project_memory.json"
    if not mf.exists():
        mf.write_text(json.dumps({"patterns": []}, ensure_ascii=False, indent=2), "utf-8")
    try:
        mem = json.loads(mf.read_text("utf-8"))
    except Exception:
        mem = {"patterns": []}
    p = {"pattern_type": a.pattern_type, "description": a.description, "importance": a.importance or "medium"}
    mem["patterns"].append(p)
    mf.write_text(json.dumps(mem, ensure_ascii=False, indent=2), "utf-8")
    print("pattern saved: " + str(p["pattern_type"]))

def main():
    p = argparse.ArgumentParser()
    p.add_argument("--project-root", default=os.getcwd())
    s = p.add_subparsers(dest="cmd")
    pi = s.add_parser("init", help="create a new project")
    pi.add_argument("--title", default="")
    pi.add_argument("--author", default="")
    pi.add_argument("--genre", default="")
    pi.set_defaults(func=cmd_init)
    s.add_parser("where").set_defaults(func=cmd_where)
    s.add_parser("project-status").set_defaults(func=cmd_status)
    pd = s.add_parser("doctor")
    pd.add_argument("--format", choices=["text", "json"], default="text")
    pd.add_argument("--chapter", default=None, help="检查指定章节")
    pd.add_argument("--deep", action="store_true", help="深度检查（含底本一致性）")
    pd.set_defaults(func=cmd_doctor)
    s.add_parser("init-contract").set_defaults(func=cmd_contract)
    pe = s.add_parser("export")
    pe.add_argument("--format", choices=["txt", "docx", "fanqie"], default="txt",
                    help="导出格式：txt 纯文本 / docx Word / fanqie 番茄平台规范")
    pe.add_argument("--output", default=None,
                    help="输出路径（目录或文件名）；不指定则写到 导出/ 下")
    pe.set_defaults(func=cmd_export)
    p6 = s.add_parser("chapter-commit")
    p6.add_argument("--chapter", required=True)
    p6.set_defaults(func=cmd_commit)
    p7 = s.add_parser("review-pipeline")
    p7.add_argument("--chapter", required=True)
    p7.add_argument("--report-file", required=True)
    p7.add_argument("--review-results")
    p7.add_argument("--metrics-out")
    p7.add_argument("--save-metrics", action="store_true")
    p7.set_defaults(func=cmd_review)
    p8 = s.add_parser("project-memory")
    p8.add_argument("action", choices=["add-pattern"])
    p8.add_argument("--pattern-type", required=True)
    p8.add_argument("--description", required=True)
    p8.add_argument("--importance", default="medium")
    p8.set_defaults(func=cmd_memory)
    a = p.parse_args()
    if hasattr(a, "func"):
        a.func(a)
    else:
        p.print_help()

if __name__ == "__main__":
    main()